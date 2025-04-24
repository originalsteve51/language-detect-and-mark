import langid
import re
import sys

from docx import Document
from docx.shared import RGBColor 

# Global values set from the command line
MAX_WORDS = None
SPECIFIED_LANG = None

def is_specified_language(text):
    """Detect if given text is in a specified language."""
    global SPECIFIED_LANG
    language, confidence = langid.classify(text)
    return language == SPECIFIED_LANG and confidence > -900.0 

def split_into_sub_phrases(phrase, max_words):
    """Split a long phrase into sub-phrases of at most max_words."""
    words = phrase.split()
    return [' '.join(words[i:i + max_words])\
                for i in range(0, len(words), max_words)]

def identify_specified_lang_phrases(old_run, para):
    global MAX_WORDS
    
    # Split the text into sentences
    sentences = re.split(r'(?<=[.!?]) +', old_run.text)  # Split by sentence endings

    specified_lang_phrases = []
    for sentence in sentences:
        # Split into sub-phrases of at most MAX_WORDS words
        sub_phrases = split_into_sub_phrases(sentence, MAX_WORDS)
        finished = False
        while not finished:
            # if not is_specified_language(sub_phrases[0]):
            if len(sub_phrases)>0:
                new_run = para.add_run(sub_phrases[0] + ' ')
                if is_specified_language(sub_phrases[0]):
                    specified_lang_phrases.append(sub_phrases[0])
                    new_run.font.color.rgb = RGBColor(255,0,255)
                

            if len(sub_phrases)>1:
                sub_phrases = sub_phrases[1::]
            else:
                finished = True
                    
    old_run.clear()

    # Useful for debug...Return what we think is highlighted
    return specified_lang_phrases

def main():
    global SPECIFIED_LANG
    global MAX_WORDS

    try:
        if len(sys.argv) == 3:
            SPECIFIED_LANG = sys.argv[1]
            MAX_WORDS = int(sys.argv[2])
            print(f'Scanning for {SPECIFIED_LANG} using a phrase length of {MAX_WORDS}')
        else:
            print('You need to specify a language code (eg de) followed by an integer phrase length')
            sys.exit()
    except ValueError:
        print('The second argument must be an integer value for the phrase length to use.')
        sys.exit()

    # Load the docx file (names are hard-coded for now)
    in_filepath = './sample_tagging_text_all_black.docx' 
    out_filepath = './sample_tagging_text_markup.docx'

    doc = Document(in_filepath)

    for para in doc.paragraphs:
        for run in para.runs:

            specified_lang_phrases = identify_specified_lang_phrases(run, para)
            
            # Following is useful for debug, see what we think was marked in the document
            if len(specified_lang_phrases) > 0:
                print(f"{len(specified_lang_phrases)} {SPECIFIED_LANG} Phrases Found:")
                
                for phrase in specified_lang_phrases:
                    print(phrase)

    # Save the document that is now decorated with language colors
    doc.save(out_filepath) 

if __name__ == "__main__":
    main()