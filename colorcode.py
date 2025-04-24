from docx import Document

from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor 

from langdetect import detect
from langdetect.lang_detect_exception import LangDetectException
from langdetect import DetectorFactory
from langdetect import detect_langs

import os
 
DetectorFactory.seed = 0

def highlight_foreign_words(docx_path, output_path):
    doc = Document(docx_path)
 
    for para in doc.paragraphs:

        '''
        for run in para.runs:
            try:
                languages = detect_langs(run.text)

                for language in languages:
                    print('assessment: ', language.lang, language.prob)
            except LangDetectException:
                pass
        '''

        for run in para.runs:

            print(run.text)

            try:
                languages = detect_langs(run.text)

                for language in languages:
                    print('assessment: ', language.lang, language.prob)
            except LangDetectException:
                pass


            words = run.text.split()

            # Make multi-word phrases
            words = words[0:9]

            for word in words:
                try:
                    lang = detect(word[0:4])
                except LangDetectException:
                    print(f'{word}: Exception')   
                    lang = 'en'  # default to English if detection fails

                new_run = para.add_run(word + ' ')
                print(f'word : {word}, {lang}')
                if lang != 'en':
                    # new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    new_run.font.color.rgb = RGBColor(255,0,255)
     
                run.clear()  # Remove original run content
 
    doc.save(output_path)
 
# Example usage:
# highlight_foreign_words("input.docx", "output_highlighted.docx")

if __name__ == "__main__":
    highlight_foreign_words("sample_tagging_text_all black.docx", "output_highlighted.docx")